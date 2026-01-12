from django.db.models import Q
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from .models import *
from django.conf import settings
from django.template.loader import render_to_string
from django.core.mail import EmailMessage
from datetime import datetime
from django.core.files.uploadedfile import UploadedFile
import fitz 
import docx 
from typing import Optional
import io


def send_templated_email(template_name, subject, recipient_list, context, attachments=None):
    context['current_year'] = datetime.now().year
    
    html_content = render_to_string(template_name, context)
    
    email = EmailMessage(
        subject,
        html_content,
        settings.DEFAULT_FROM_EMAIL,
        recipient_list
    )
    
    email.content_subtype = "html" 
    
    if attachments:
        for filename, content, mimetype in attachments:
            email.attach(filename, content, mimetype)
    
    try:
        email.send()
        return True
    except Exception as e:
        import traceback
        print(f"Error sending email: {e}\n{traceback.format_exc()}")
        return False

def get_subscribers_context(request):
    """
    Helper function to get course enrollment list and pagination context.
    Matches the 'students' variable used in the template.
    """
    query = request.GET.get('q', '')
    
    enrollment_queryset = CoursePurchase.objects.filter(status='success').select_related('user').order_by('-created_at')
    
    if query:
        enrollment_queryset = enrollment_queryset.filter(
            Q(user__email__icontains=query) | 
            Q(user__first_name__icontains=query) | 
            Q(user__last_name__icontains=query) |
            Q(course_name__icontains=query)
        )
    
    paginator = Paginator(enrollment_queryset, 10)
    page_number = request.GET.get('page')

    try:
        students = paginator.page(page_number)
    except PageNotAnInteger:
        students = paginator.page(1)
    except EmptyPage:
        students = paginator.page(paginator.num_pages)

    return {
        'students': students,
        'query': query,
    }


def extract_resume_text(file: UploadedFile) -> Optional[str]:
    """
    Extracts readable text from an uploaded resume file (PDF or DOCX) in-memory.
    - Accepts a Django UploadedFile (InMemoryUploadedFile or TemporaryUploadedFile).
    - Returns the plain text or None if extraction fails.
    """
    try:
        # Determine file extension
        ext = file.name.lower().split('.')[-1]

        # --- PDF Extraction ---
        if ext == "pdf":
            # Read bytes into PyMuPDF
            file_bytes = file.read()
            pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_parts = [page.get_text("text") for page in pdf_doc if page.get_text("text").strip()]
            pdf_doc.close()
            final_text = "\n".join(text_parts).strip()
            return final_text if final_text else None

        # --- DOCX Extraction ---
        elif ext == "docx":
            file_bytes = io.BytesIO(file.read())
            doc = docx.Document(file_bytes)
            full_text = [para.text for para in doc.paragraphs if para.text.strip()]

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

            final_text = "\n".join(full_text).strip()
            return final_text if final_text else None

        # --- Unsupported formats ---
        elif ext == "doc":
            print("Extraction Warning: Legacy .doc format detected. Conversion required.")
            return None
        else:
            print(f"Extraction Error: Unsupported file extension .{ext}")
            return None

    except Exception as e:
        print(f"CRITICAL ERROR during text extraction: {str(e)}")
        return None